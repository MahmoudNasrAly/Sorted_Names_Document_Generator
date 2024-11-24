from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_table_borders(table):
    # Loop through each cell in the table
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Single line border
                border.set(qn('w:sz'), '12')  # Border size, thicker for visibility
                border.set(qn('w:color'), '000000')  # Black color
                tcBorders.append(border)

            tcPr.append(tcBorders)


def sort_arabic_records(records, file_name):
    # Sort the records in ascending order
    sorted_records = sorted(records)

    # Create a Word document
    doc = Document()

    # Number of columns: 1 for ID, 1 for the name, and 15 empty columns
    num_cols = 17

    # Add a table to the document with rows for each record
    table = doc.add_table(rows=len(sorted_records), cols=num_cols)

    # Set table borders to be thick and black
    set_table_borders(table)

    # Populate the table with IDs and names, align names to the right
    for i, record in enumerate(sorted_records):
        # Fill ID column (first column)
        table.cell(i, 0).text = str(i + 1)

        # Fill name column (second column) and align to the right
        name_cell = table.cell(i, 1)
        name_cell.text = record
        name_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save the document with the given file name
    doc.save(f'{file_name}.docx')


if __name__ == "__main__":
    records = []

    while True:
        # Get the record from the user
        record = input("Enter a name record (or type 'done' to finish): ")

        # Check if the user wants to finish inputting records
        if record.lower() == 'done':
            break

        # Add the record to the list
        records.append(record)

    # Ask the user for the file name
    file_name = input("Enter the name for the Word file (without extension): ")

    # Sort the records and save them to a Word file with the new name
    sort_arabic_records(records, file_name)

    print(f"Records have been sorted and saved in '{file_name}.docx'.")
